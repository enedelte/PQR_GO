import re
from docx import Document
import os

directorio = "C:/Users/daale/OneDrive - NUEVA EPS/GO_PQR/GESTIÓN_PQR/MAYO"
#document = Document("VO-GA-DGO-2445684-23.docx")

ruta_salida = 'C:/Users/daale/OneDrive - NUEVA EPS/GO_PQR/'

email_info = r'/b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+/.[A-Za-z]{2,}/b'

# Lista para almacenar las direcciones de correo electrónico encontradas
email_list = []

# Recorrer los archivos en el directorio
for archivo in os.listdir(directorio):
    if archivo.endswith('.docx'):
        # Construir la ruta completa del archivo
        ruta_archivo = os.path.join(directorio, archivo)
        
        # Cargar el documento
        document = Document(ruta_archivo)
        
        # Buscar direcciones de correo electrónico en los párrafos del documento
        for paragraph in document.paragraphs:
            matches = re.findall(email_info, paragraph.text)
            if matches:
                email_list.extend(matches)

# Imprimir la lista de direcciones de correo electrónico encontradas
for email in email_list:
    print(email_info)

